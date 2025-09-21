"""
file_profiler_agent.py

A focused file profiler for analyzing single files. It uses PostgreSQL for tracking,
extracts text from various document types, and leverages Google's Gemini model via LangChain for analysis.

This version is streamlined to focus on the single-file analysis function.

Usage (as a library):
    from file_profiler_pg_gemini import FileAnalyzer
    analyzer = FileAnalyzer()
    result = analyzer.analyze_single_file('/path/to/your/file.pdf')
    print(result)

Usage (from command line):
    python file_profiler_pg_gemini.py --file-path /path/to/your/file.pdf

Requirements:
    pip install pymupdf pandas python-docx psycopg2-binary python-dotenv langchain langchain-core langchain-google-genai google-generativeai
"""
import os
import sys
import time
import json
import argparse
import hashlib
import traceback
from pathlib import Path

# Third-party libraries for file processing
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    import pandas as pd
except ImportError:
    pd = None

try:
    import docx
except ImportError:
    docx = None

# Postgres
import psycopg2
from psycopg2.pool import ThreadedConnectionPool
import base64

# LangChain with Google Gemini
try:
    from langchain_google_genai.chat_models import ChatGoogleGenerativeAI
    from langchain_core.prompts import ChatPromptTemplate
    from langchain_core.output_parsers import JsonOutputParser
    from google.api_core import exceptions as google_exceptions
except ImportError:
    ChatGoogleGenerativeAI = None
    google_exceptions = None
try:
    # Client for the new PDF extraction method
    from mistralai import Mistral
except ImportError:
    Mistral = None


from dotenv import load_dotenv
load_dotenv()
# ----------------- CONFIG -----------------
DB_HOST = os.getenv("PG_HOST", "localhost")
DB_PORT = int(os.getenv("PG_PORT", 5432))
DB_NAME = os.getenv("PG_DB", "file_analysis")
DB_USER = os.getenv("PG_USER", "postgres")
DB_PASS = os.getenv("PG_PASS", "admin")
DB_MIN_POOL = int(os.getenv("PG_POOL_MIN", 1))
DB_MAX_POOL = int(os.getenv("PG_POOL_MAX", 8))

# Switched to Google Gemini
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
MODEL_NAME = os.getenv("MODEL_NAME", "gemini-2.5-flash-lite")
MISTRAL_API_KEY = os.getenv("MISTRAL_OCR_KEY")

DEFAULT_MAX_RPM = 3
DEFAULT_TEXT_LIMIT = 8000 # Increased limit for better context
MAX_RETRIES = 5
INITIAL_BACKOFF = 8.0

SUPPORTED_EXTS = {".pdf", ".xlsx", ".xls", ".xlsm", ".docx"}

# ----------------- FILE ANALYZER CLASS -----------------

class FileAnalyzer:
    """Encapsulates file profiling logic."""

    def __init__(self):
        """Initializes the analyzer, DB pool, and LangChain LLM."""
        self._conn_pool = None
        self.init_db()
        if ChatGoogleGenerativeAI and GOOGLE_API_KEY:
            self.llm = ChatGoogleGenerativeAI(model=MODEL_NAME, google_api_key=GOOGLE_API_KEY)
        else:
            self.llm = None
        if Mistral and MISTRAL_API_KEY:
            self.ocr_client = Mistral(api_key=MISTRAL_API_KEY)
        else:
            self.ocr_client = None

        self._last_call_ts = 0.0

    def get_db_pool(self):
        """Creates and returns a threaded database connection pool."""
        if self._conn_pool is None:
            self._conn_pool = ThreadedConnectionPool(
                minconn=DB_MIN_POOL, maxconn=DB_MAX_POOL,
                host=DB_HOST, port=DB_PORT, dbname=DB_NAME, user=DB_USER, password=DB_PASS
            )
        return self._conn_pool

    def init_db(self):
        """Initializes the database schema if it doesn't exist."""
        conn = None
        try:
            conn = psycopg2.connect(host=DB_HOST, port=DB_PORT, dbname=DB_NAME, user=DB_USER, password=DB_PASS)
            cur = conn.cursor()
            cur.execute("""
                CREATE TABLE IF NOT EXISTS files (
                    id BIGSERIAL PRIMARY KEY,
                    file_path TEXT UNIQUE NOT NULL,
                    file_hash TEXT,
                    file_size BIGINT,
                    extract_status TEXT DEFAULT 'pending', -- pending / done / error
                    extract_error TEXT,
                    extracted_text TEXT,
                    analyzed BOOLEAN DEFAULT FALSE,
                    analysis_json JSONB,
                    analysis_error TEXT,
                    updated_at TIMESTAMP DEFAULT NOW()
                );
            """)
            cur.execute("CREATE INDEX IF NOT EXISTS idx_files_hash ON files(file_hash);")
            conn.commit()
            cur.close()
        finally:
            if conn:
                conn.close()

    def _compute_file_hash(self, path: Path, chunk_size: int = 65536) -> str:
        """Computes the SHA256 hash of a file."""
        h = hashlib.sha256()
        with path.open("rb") as f:
            for chunk in iter(lambda: f.read(chunk_size), b""):
                h.update(chunk)
        return h.hexdigest()

    def _extract_text(self, path: Path) -> (str, str):
        """Extracts text from a file, returns (text, error_string)."""
        try:
            ext = path.suffix.lower()
            if ext == ".pdf":
                txt = self._extract_pdf_text(path)
            elif ext in {".xlsx", ".xls", ".xlsm"}:
                txt = self._extract_excel_text(path)
            elif ext == ".docx":
                txt = self._extract_docx_text(path)
            else:
                txt = "[Unsupported file type]"

            if txt.startswith("[") and "Error" in txt:
                return "", txt
            return txt, ""
        except Exception:
            return "", f"Unexpected extraction error: {traceback.format_exc()}"

    def _extract_pdf_text(self, path: Path) -> str:
        """
        Creates a smaller in-memory PDF with the first 20 pages and sends that for OCR
        to avoid request size limits.
        """
        if not self.ocr_client:
            return "[PDF extraction error: Mistral client not initialized. Check MISTRAL_API_KEY]"
        if fitz is None:
            return "[PDF processing error: PyMuPDF (fitz) is required to select pages but is not installed]"
            
        try:
            # Step 1: Create a new PDF in memory with only the first 20 pages
            pdf_bytes_to_send = b''
            with fitz.open(path) as original_doc:
                if original_doc.page_count == 0:
                    return "" # Handle empty PDF

                # Determine how many pages to copy (up to 20)
                page_count_to_extract = min(original_doc.page_count, 20)
                
                with fitz.open() as new_doc:
                    new_doc.insert_pdf(original_doc, from_page=0, to_page=page_count_to_extract - 1)
                    pdf_bytes_to_send = new_doc.tobytes()

            # Step 2: Base64 encode the smaller PDF's bytes
            encoded_pdf = base64.b64encode(pdf_bytes_to_send).decode("utf-8")
            
            # Step 3: Send the smaller payload to the OCR service
            document = {"type": "document_url", "document_url": f"data:application/pdf;base64,{encoded_pdf}"}
            ocr_response = self.ocr_client.ocr.process(model="mistral-ocr-latest", document=document, include_image_base64=False)
            time.sleep(1)
            
            pages = ocr_response.pages if hasattr(ocr_response, "pages") else ocr_response
            
            # Step 4: Combine the markdown from the received pages
            page_texts = [f"--- Page {i+1} ---\n{page.markdown}" for i, page in enumerate(pages)]
            return "\n\n".join(page_texts).strip()

        except Exception as e:
            return f"[PDF Extraction Error (OCR): {e}]"

    def _extract_excel_text(self, path: Path) -> str:
        if pd is None:
            return "[Excel extraction error: pandas not installed]"
        try:
            # Using string dtype to prevent pandas from auto-formatting numbers
            sheets = pd.read_excel(path, sheet_name=None, dtype=str)
            parts = [f"--- Sheet: {name} ---\n{df.to_csv(index=False)}" for name, df in sheets.items()]
            return "\n".join(parts).strip()
        except Exception as e:
            return f"[Excel Extraction Error: {e}]"

    def _extract_docx_text(self, path: Path) -> str:
        if docx is None:
            return "[DOCX extraction error: python-docx not installed]"
        try:
            doc = docx.Document(str(path))
            return "\n".join(p.text for p in doc.paragraphs).strip()
        except Exception as e:
            return f"[DOCX Extraction Error: {e}]"

    def _throttle(self, max_rpm: int):
        """Ensures the rate of API calls does not exceed max_rpm."""
        min_gap = 60.0 / max(1, max_rpm)
        now = time.time()
        elapsed = now - self._last_call_ts
        if elapsed < min_gap:
            time.sleep(min_gap - elapsed)
        self._last_call_ts = time.time()

    def _classify_with_llm(self, text: str, file_name: str, text_limit: int):
        """Sends text to Gemini via LangChain for classification and returns a JSON object."""
        if not self.llm:
            return {"ok": False, "error": "Gemini client not initialized. Check GOOGLE_API_KEY."}

        prompt_template = ChatPromptTemplate.from_template("""
            You are an expert file diagnoser for economic reports. Your task is to analyze the provided text from a document and return ONLY a compact JSON object.

            ## JSON Keys & Instructions:

            - **file_name**: The name of the file.
            - **domain**: Identify the **specific economic report name or key indicator** from the document's content itself. The value must be derived from the text.
                - *Examples for format*: "Consumer Price Index", "Index of Industrial Production", "State Finances Report", "Foreign Trade Statistics".
            - **subdomain**: A short tag classifying the document type.
                - *Examples for format*: "Press Release", "Committee Minutes", "Annual Report", "Survey Data".
            - **intents**: Based on the text, state the **primary querying purpose** of this file as a concise, action-oriented phrase.
                - *Examples for format*: "To query monthly inflation rates," "To analyze industrial growth sector-wise," "To track import-export data," "To understand trends in state government debt."
            - **publishing_authority**: The official organization that published the document.
            - **published_date**: The publication month and year in strictly in "Month YYYY" format, if it is not present return the start of reference date in "Month YYYY" format.
            - **period_of_reference**: The time period the data in the document refers to.
            - **brief_summary**: A concise, neutral summary of the document's purpose (max 3 sentences).
            - **quality_score**: A score from 1-3 on how structured and parsable the document text is.

            ---
            ## Critical Rules:

            1.  Output **only** the raw JSON object and nothing else.
            2.  The examples provided are for **guidance on the expected format only**.
            3.  You **MUST** determine the `domain` and `intents` by analyzing the actual file content. **DO NOT simply copy the examples.**

            ---

            **File**: {file_name}
            **Content**:
            {text}
            """.strip())

        parser = JsonOutputParser()
        chain = prompt_template | self.llm | parser
        
        backoff = INITIAL_BACKOFF
        for attempt in range(1, MAX_RETRIES + 1):
            try:
                content = chain.invoke({
                    "file_name": file_name,
                    "text": text[:text_limit]
                })
                return {"ok": True, "json": content}
            except (google_exceptions.ResourceExhausted, google_exceptions.ServiceUnavailable, google_exceptions.DeadlineExceeded) as e:
                print(f"API Error (attempt {attempt}/{MAX_RETRIES}): {e}. Retrying in {backoff}s...")
                if attempt < MAX_RETRIES:
                    time.sleep(backoff)
                    backoff *= 2
                else:
                    return {"ok": False, "error": f"Google API error after retries: {e}", "traceback": traceback.format_exc()}
            except Exception as e:
                # This catches parsing errors from JsonOutputParser or other unexpected issues
                return {"ok": False, "error": str(e), "traceback": traceback.format_exc()}

    def analyze_single_file(self, file_path: str, text_limit: int = DEFAULT_TEXT_LIMIT, max_rpm: int = DEFAULT_MAX_RPM) -> dict:
        """
        Analyzes a single file by extracting its text and sending it to an LLM.

        Args:
            file_path: The absolute path to the file.
            text_limit: The maximum number of characters to send for analysis.
            max_rpm: The maximum requests per minute for the LLM API.

        Returns:
            A dictionary containing the analysis result or an error message.
        """
        p = Path(file_path)
        if not p.is_file() or p.suffix.lower() not in SUPPORTED_EXTS:
            return {"error": "File does not exist or is not a supported type.", "file_path": file_path}

        # 1. Extract text
        print(f"Extracting text from {file_path}")
        extracted_text, err = self._extract_text(p)

        if err:
            print(f"Extraction failed: {err}")
            return {"error": f"Extraction failed: {err}", "file_path": file_path}

        # 2. Throttle and Analyze
        print(f"Analyzing {file_path} using {MODEL_NAME}...")
        self._throttle(max_rpm)
        analysis_result = self._classify_with_llm(extracted_text, p.name, text_limit)

        # 3. Store results in DB for tracking
        conn = self.get_db_pool().getconn()
        try:
            with conn.cursor() as cur:
                file_hash = self._compute_file_hash(p)
                file_size = p.stat().st_size
                if analysis_result.get("ok"):
                    cur.execute("""
                        INSERT INTO files (file_path, file_hash, file_size, extract_status, extracted_text, analyzed, analysis_json, updated_at)
                        VALUES (%s, %s, %s, 'done', %s, TRUE, %s, NOW())
                        ON CONFLICT (file_path) DO UPDATE SET
                            file_hash = EXCLUDED.file_hash, file_size = EXCLUDED.file_size,
                            extract_status = 'done', extracted_text = EXCLUDED.extracted_text,
                            analyzed = TRUE, analysis_json = EXCLUDED.analysis_json,
                            analysis_error = NULL, updated_at = NOW();
                    """, (file_path, file_hash, file_size, extracted_text, json.dumps(analysis_result['json'])))
                else:
                    cur.execute("""
                        INSERT INTO files (file_path, file_hash, file_size, extract_status, extracted_text, analyzed, analysis_error, updated_at)
                        VALUES (%s, %s, %s, 'done', %s, FALSE, %s, NOW())
                        ON CONFLICT (file_path) DO UPDATE SET
                            file_hash = EXCLUDED.file_hash, file_size = EXCLUDED.file_size,
                            extract_status = 'done', extracted_text = EXCLUDED.extracted_text,
                            analyzed = FALSE, analysis_json = NULL,
                            analysis_error = EXCLUDED.analysis_error, updated_at = NOW();
                    """, (file_path, file_hash, file_size, extracted_text, json.dumps(analysis_result)))
                conn.commit()
        finally:
            self.get_db_pool().putconn(conn)

        return analysis_result

# ----------------- CLI -----------------
def make_parser():
    p = argparse.ArgumentParser(description="Analyzes a single file using Gemini and returns a JSON analysis.")
    p.add_argument("--file-path", required=True, help="Path to the file to analyze")
    p.add_argument("--max-rpm", type=int, default=DEFAULT_MAX_RPM, help="Max requests per minute for LLM API")
    p.add_argument("--text-limit", type=int, default=DEFAULT_TEXT_LIMIT, help="Max characters to send to LLM")
    return p

def main():
    parser = make_parser()
    args = parser.parse_args()
    analyzer = FileAnalyzer()
    print(f"Analyzing single file: {args.file_path}")
    result = analyzer.analyze_single_file(args.file_path, text_limit=args.text_limit, max_rpm=args.max_rpm)
    print(json.dumps(result, indent=2))

if __name__ == "__main__":
    if len(sys.argv) == 1:
        print("This script is intended to be used with command-line arguments.")
        print("Example: python your_script_name.py --file-path /path/to/document.pdf")
        if os.getenv("GOOGLE_API_KEY"):
            print("\nRunning a test with a dummy file...")
            analyzer_instance = FileAnalyzer()
            dummy_file_path = "dummy_document_for_analysis.docx"
            try:
                if not Path(dummy_file_path).exists():
                    print(f"Creating a dummy file: {dummy_file_path}")
                    doc = docx.Document()
                    doc.add_heading('Ministry of Finance', 0)
                    doc.add_paragraph('Press Release: Quarterly Economic Outlook')
                    doc.add_paragraph('Date: July 2025')
                    doc.add_paragraph('This report covers the economic performance for the first quarter of 2025 (Q1 2025).')
                    doc.save(dummy_file_path)
                
                analysis = analyzer_instance.analyze_single_file(dummy_file_path)
                print("\n--- Analysis Result ---")
                print(json.dumps(analysis, indent=2))
                # Clean up the dummy file
                os.remove(dummy_file_path)

            except Exception as e:
                print(f"Could not create or analyze dummy file. Error: {e}")
        else:
            print("\nGOOGLE_API_KEY environment variable not set. Cannot run example.")
    else:
        main()